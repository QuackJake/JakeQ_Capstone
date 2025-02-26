import pypdf, os, json, docx, re, sys
from lxml import etree
from docx import Document
from typing import Union, List, Dict
from pypdf.errors import EmptyFileError
from dataclasses import dataclass
from pypdf import PdfReader, PdfWriter
from datetime import datetime

class pdf_reader:
    def __init__(self, target_dir: str):
        self.target_dir = target_dir
        if not os.path.exists(target_dir):
            raise ValueError(f"Target directory does not exist: {target_dir}")
        
    def exception_wrapper(wrapped_function):
        def _wrapper(*args, **kwargs):
            try:
                result = wrapped_function(*args, **kwargs)
                return result
            except FileNotFoundError as e:
                print(f"File error: {e}")
            except EmptyFileError as e:
                print(f"Empty file error: {e}")
            except Exception as e:
                print(f"An error occurred: {e}")
            
        return _wrapper

    @exception_wrapper
    def pdf_diagnostic(self, filename):
        file_path = os.path.join(self.target_dir, filename)
        print(f"\nDiagnostic checks for: {file_path}")
        print(f"1. Full absolute path: {os.path.abspath(file_path)}")
        print(f"2. File exists: {os.path.exists(file_path)}")
        if os.path.exists(file_path):
            print(f"3. File size: {os.path.getsize(file_path)} bytes")
            print(f"4. Is file readable: {os.access(file_path, os.R_OK)}")
            print(f"5. Parent directory exists: {os.path.exists(os.path.dirname(file_path))}")
            
            print("\nMetadata:")
            self.read_metadata(filename)
        else:
            print("3-5. Skipped as file doesn't exist")

    def _get_pdf_path(self, filename: str) -> str:
        full_path = os.path.join(self.target_dir, filename)
        if not os.path.exists(full_path):
            raise FileNotFoundError(f"PDF file not found: {full_path}")
        if os.path.getsize(full_path) == 0:
            raise EmptyFileError(f"PDF file is empty: {full_path}")
        return full_path

    def _get_output_path(self, filename: str) -> str:
        output_path = os.path.join(self.target_dir, filename)
        output_dir = os.path.dirname(output_path)
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        return output_path

    def get_page_count(self, filename: str) -> int:
        try:
            pdf_path = self._get_pdf_path(filename)
            reader = pypdf.PdfReader(pdf_path)
            return len(reader.pages)
        except EmptyFileError as e:
            raise EmptyFileError(f"Cannot read empty PDF file: {filename}") from e
        except Exception as e:
            raise RuntimeError(f"Error reading PDF file {filename}: {str(e)}") from e

    @exception_wrapper
    def extract_text_from_pages(self, filename: str, all_pages: bool = True, page_count: int = None) -> List[str]:
        if not all_pages and page_count is None:
            raise ValueError("page_count is required when all_pages is False")
            
        try:
            pdf_path = self._get_pdf_path(filename)
            reader = pypdf.PdfReader(pdf_path)
            extracted_text = []
            
            for i, page in enumerate(reader.pages):
                if not all_pages and i >= page_count:
                    break
                text = page.extract_text()
                extracted_text.append(text)
                
            return extracted_text
        except Exception as e:
            raise RuntimeError(f"Error extracting text from PDF {filename}: {str(e)}") from e

    @exception_wrapper
    def extract_bytes_from_pages(self, filename: str, all_pages: bool = True, page_count: int = None) -> List[Union[bytes, None]]:
        """Extract raw bytes content from PDF pages."""
        if not all_pages and page_count is None:
            raise ValueError("page_count is required when all_pages is False")
            
        try:
            pdf_path = self._get_pdf_path(filename)
            reader = pypdf.PdfReader(pdf_path)
            extracted_bytes = []
            
            for i, page in enumerate(reader.pages):
                if not all_pages and i >= page_count:
                    break
                content = page.get_contents()
                extracted_bytes.append(content)
                
            return extracted_bytes
        except Exception as e:
            raise RuntimeError(f"Error extracting bytes from PDF {filename}: {str(e)}") from e

    @exception_wrapper
    def write_text_to_file(self, pdf_filename: str, output_filename: str = None, all_pages: bool = True, page_count: int = None) -> str:
        if  not output_filename:
            output_filename = (pdf_filename.split('.')[0]) + '.txt'
        try:
            pages = self.extract_text_from_pages(pdf_filename, all_pages, page_count)
            output_path = self._get_output_path(output_filename)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(f"Extracted from: {pdf_filename}\n")
                f.write(f"Total pages: {self.get_page_count(pdf_filename)}\n\n")
                
                for i, page_content in enumerate(pages, 1):
                    f.write(f"--- Page {i} ---\n")
                    f.write(page_content)
                    f.write("\n\n")
            
            return output_path
        except Exception as e:
            raise RuntimeError(f"Error writing text to file {output_filename}: {str(e)}") from e

    @exception_wrapper
    def write_bytes_to_file(self, pdf_filename: str, output_filename: str, all_pages: bool = True, page_count: int = None) -> str:
        """Write raw bytes content to a file."""
        try:
            pages = self.extract_bytes_from_pages(pdf_filename, all_pages, page_count)
            output_path = self._get_output_path(output_filename)
            
            with open(output_path, 'wb') as f:
                f.write(f"Extracted from: {pdf_filename}\n".encode('utf-8'))
                f.write(f"Total pages: {self.get_page_count(pdf_filename)}\n\n".encode('utf-8'))
                
                for i, page_content in enumerate(pages, 1):
                    f.write(f"--- Page {i} ---\n".encode('utf-8'))
                    if page_content is not None:
                        f.write(page_content)
                    else:
                        f.write(b"[No content available for this page]")
                    f.write(b"\n\n")
            
            return output_path
        except Exception as e:
            raise RuntimeError(f"Error writing bytes to file {output_filename}: {str(e)}") from e
    
    @exception_wrapper
    def parse_clio_pdf(self, filename: str, all_pages: bool = True, page_count: int = None):
        test = self.extract_text_from_pages(filename, all_pages, page_count)[0]
        st = test.split('\n')

        keys = ['export_date', 'export_time', 'name', 'dob', 'phone', 'email', 'entity', 'client']
        values = st[3].split(' ')
        values[0:3] = [' '.join(values[0:3])]
        values.insert(0, st[1])
        values.insert(0, st[0].split(' ')[2])
        # page_count = st[4]

        data = dict(zip(keys, values))
        json_object = json.dumps(data, indent=4)
        print(json_object)
        # with open("output.json", "w") as json_file:
        #     json.dump(data, json_file, indent=4)
        # print(keys)
        # print(values)

        # Test Contacts Format:
        # 'Contacts Export 01/08/2025' > Date of Export
        # '10:50 PM' > Time of Export
        # 'Name Date of Birth Primary Phone Number Primary Email Address Entity Client' > Header Column
        # 'John Test Smith 01/01/2025 1111111111 johntest@test.com Person True' > Values
        # '1/1' > PageNum

        utc_time = "-05'00'"  # UTC time optional
    
    @exception_wrapper
    def create_metadata(self, metadata_dict=None):
        default_metadata = {
            "/Author": "Unknown Author",
            "/Producer": "Unknown Producer",
            "/Title": "Untitled",
            "/Subject": "No Subject",
            "/Keywords": "",
            "/CreationDate": datetime.now(),
            "/ModDate": datetime.now(),
            "/Creator": "Unknown Creator",
            "/CustomField": ""
        }
        if metadata_dict:
            default_metadata.update(metadata_dict)
        return default_metadata

    @exception_wrapper
    def update_metadata(self, filename: str, metadata_dict=None):
        file_path = os.path.join(self.target_dir, filename)
        reader = PdfReader(file_path)
        writer = PdfWriter(clone_from=file_path)

        metadata = self.create_metadata(metadata_dict)
        print(metadata)
        
        writer.add_metadata(metadata)
        
        with open(f"C:/Capstone/Capstone_2025/pdf_reader_src/legal_documents/pdf/{filename}", "wb") as f:
            writer.write(f)
    
    @exception_wrapper
    def read_metadata(self, filename: str):
        file_path = os.path.join(self.target_dir, filename)
        meta = PdfReader(file_path).metadata
        for key in ["/Title", "/Author", "/Subject", "/Creator", "/Producer", "/CreationDate", "/ModDate"]:
            print(f"{key}: {meta.get(key, 'Not Available')}")

    @exception_wrapper
    def write_metadata(self, filename: str, metadata_dict=None):
        file_path = os.path.join(self.target_dir, filename)
        reader = PdfReader(file_path)
        writer = PdfWriter()
        
        for page in reader.pages:
            writer.add_page(page)
        
        writer.add_metadata(reader.metadata or {})
        writer.add_metadata(self.create_metadata(metadata_dict))
        
        with open(file_path, "wb") as f:
            writer.write(f)

class docx_reader:
    def __init__(self, target_dir: str):
        self.target_dir = target_dir
        if not os.path.exists(target_dir):
            raise ValueError(f"Target directory does not exist: {target_dir}")
        
    def exception_wrapper(wrapped_function):
        def _wrapper(*args, **kwargs):
            try:
                result = wrapped_function(*args, **kwargs)
                return result
            except FileNotFoundError as e:
                print(f"File error: {e}")
            except EmptyFileError as e:
                print(f"Empty file error: {e}")
            except Exception as e:
                print(f"An error occurred: {e}")
        return _wrapper
    
    @exception_wrapper
    def get_text(self, filename):
        """
        Extract text from a Word document including various field types and content controls.
        Handles Unicode characters properly.
        
        Args:
            filename (str): Path to the Word document
            
        Returns:
            str: Extracted text with fields and controls
        """
        doc = Document(filename)
        fullText = []
    
        def clean_text(text):
            """Clean and normalize text to handle Unicode characters"""
            if text is None:
                return ""
            # Replace problematic characters or normalize them
            text = text.replace('\u25ba', '->') # Replace right-pointing triangle
            text = text.replace('\u2022', '*')  # Replace bullet point
            text = text.replace('\u2013', '-')  # Replace en dash
            text = text.replace('\u2014', '--') # Replace em dash
            text = text.replace('\u2018', "'")  # Replace left single quote
            text = text.replace('\u2019', "'")  # Replace right single quote
            text = text.replace('\u201c', '"')  # Replace left double quote
            text = text.replace('\u201d', '"')  # Replace right double quote
            return text.strip()
        
        # Get regular paragraph text
        for para in doc.paragraphs:
            if para.text.strip():
                fullText.append(clean_text(para.text))
        
        # Get table content
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        fullText.append(clean_text(cell.text))
        
        # Parse XML for additional elements
        docx_xml = doc._element.xml
        root = etree.fromstring(docx_xml)
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'v': 'urn:schemas-microsoft-com:vml',
            'w14': 'http://schemas.microsoft.com/office/word/2010/wordml'
        }
        
        # Get content controls (structured document tags)
        for sdt in root.findall('.//w:sdt', namespaces):
            try:
                # Get the tag name/title if available
                tag = sdt.find('.//w:tag', namespaces)
                tag_val = tag.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') if tag is not None else ''
                
                # Get all text elements within the control
                text_elements = sdt.findall('.//w:t', namespaces)
                control_text = ''.join([clean_text(el.text) for el in text_elements if el.text])
                
                if control_text:
                    if tag_val:
                        fullText.append(f"{clean_text(tag_val)}: {control_text}")
                    else:
                        fullText.append(control_text)
            except Exception as e:
                print(f"Error processing SDT: {str(e)}", file=sys.stderr)
        
        # Get simple fields
        for field in root.findall('.//w:fldSimple', namespaces):
            try:
                # Get the field instruction
                instr = field.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instr', '')
                field_text = field.find('.//w:t', namespaces)
                
                if field_text is not None and field_text.text:
                    if instr:
                        fullText.append(f"{clean_text(instr)}: {clean_text(field_text.text)}")
                    else:
                        fullText.append(clean_text(field_text.text))
            except Exception as e:
                print(f"Error processing simple field: {str(e)}", file=sys.stderr)
        
        # Get complex fields
        for field_char in root.findall('.//w:fldChar', namespaces):
            try:
                fld_char_type = field_char.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType')
                if fld_char_type == 'begin':
                    # Find the corresponding instruction text
                    instr_text = field_char.getparent().getnext()
                    if instr_text is not None:
                        instr = instr_text.find('.//w:instrText', namespaces)
                        if instr is not None and instr.text:
                            fullText.append(clean_text(instr.text))
            except Exception as e:
                print(f"Error processing complex field: {str(e)}", file=sys.stderr)
        
        # Get text boxes (shapes)
        for text_box in root.findall('.//v:textbox', namespaces):
            try:
                text_elements = text_box.findall('.//w:t', namespaces)
                text_box_content = ''.join([clean_text(el.text) for el in text_elements if el.text])
                if text_box_content:
                    fullText.append(text_box_content)
            except Exception as e:
                print(f"Error processing text box: {str(e)}", file=sys.stderr)
        
        # Filter out empty strings and join
        filtered_text = list(filter(None, fullText))
        
        # Write if required
        if hasattr(self, 'write'):
            try:
                self.write(filtered_text)
            except Exception as e:
                print(f"Error writing output: {str(e)}", file=sys.stderr)
        
        try:
            return '\n'.join(filtered_text)
        except UnicodeEncodeError as e:
            # If joining fails due to encoding, try encoding as UTF-8
            return '\n'.join(filtered_text).encode('utf-8', errors='replace').decode('utf-8')
    
    @exception_wrapper
    def write(self, input: list[str], filename: str = None):
        if not filename:
            filename = 'default.txt' # Default file write name
        with open(filename, 'w') as t:
            for line in input:
                t.write(line + '\n')
        msg = filename + ' successfully written'
        print(msg)

    @exception_wrapper
    def get_files(self):
        files = os.listdir(self.target_dir)
        return files
    
    @exception_wrapper
    def format_docx(self):
        # formatted_text, fulltext = self.get_text(self.target_dir + '1353FA_Certificate_of_Service_of_Financial_Declaration.docx')
        # self.write(fulltext)
        # print(fulltext)
        
        dir = self.target_dir + '1353FA_Certificate_of_Service_of_Financial_Declaration.docx'
        
        parseda, parseds, c = self.process_legal_form(dir)
        print(parseda)

    @exception_wrapper
    def extract_metadata(filename: str):
        doc = docx.Document(filename)  # Create a Document object from the Word document file.
        core_properties = doc.core_properties  # Get the core properties of the document.
        metadata = {}  # Initialize an empty dictionary to store metadata
        # Extract core properties
        for prop in dir(core_properties):  # Iterate over all properties of the core_properties object.
            if prop.startswith('__'):  # Skip properties starting with double underscores (e.g., __elenent). Not needed
                continue
            value = getattr(core_properties, prop)  # Get the value of the property.
            if callable(value):  # Skip callable properties (methods).
                continue
            if prop == 'created' or prop == 'modified' or prop == 'last_printed':  # Check for datetime properties.
                if value:
                    value = value.strftime('%Y-%m-%d %H:%M:%S')  # Convert datetime to string format.
                else:
                    value = None
            metadata[prop] = value  # Store the property and its value in the metadata dictionary.
        # Extract custom properties (if available).
        try:
            custom_properties = core_properties.custom_properties  # Get the custom properties (if available).
            if custom_properties:  # Check if custom properties exist.
                metadata['custom_properties'] = {}  # Initialize a dictionary to store custom properties.
                for prop in custom_properties:  # Iterate over custom properties.
                    metadata['custom_properties'][prop.name] = prop.value  # Store the custom property name and value.
        except AttributeError:
            # Custom properties not available in this version.
            pass  # Skip custom properties extraction if the attribute is not available.
        return metadata


    

@dataclass
class FormField:
    field_type: str  # 'checkbox' or 'text_entry'
    label: str       # The text associated with the field
    value: str       # The current value or empty string
    position: int    # Position in document for reference

class LegalFormParser:
    def __init__(self):
        self.checkbox_pattern = r'\[\s*[Xx]?\s*\]'  # Matches [  ] or [X]
        self.blank_line_pattern = r'_{3,}'          # Matches 3 or more underscores
        
    def extract_fields(self, filename: str) -> Dict[str, List[FormField]]:
        """
        Extract form fields from a legal document.
        
        Args:
            filename: Path to the Word document
            
        Returns:
            Dictionary containing lists of checkboxes and text entry fields
        """
        doc = Document(filename)
        fields = {
            'checkboxes': [],
            'text_entries': []
        }
        
        position = 0
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
                
            # Find checkbox fields
            checkbox_matches = list(re.finditer(self.checkbox_pattern, text))
            for match in checkbox_matches:
                # Get the label by looking at text after the checkbox
                start_pos = match.end()
                # Find the next checkbox or end of line
                next_checkbox = re.search(self.checkbox_pattern, text[start_pos:])
                end_pos = start_pos + next_checkbox.start() if next_checkbox else len(text)
                label = text[start_pos:end_pos].strip()
                
                fields['checkboxes'].append(FormField(
                    field_type='checkbox',
                    label=label,
                    value='checked' if 'X' in match.group() else 'unchecked',
                    position=position
                ))
            
            # Find text entry fields
            blank_matches = list(re.finditer(self.blank_line_pattern, text))
            for match in blank_matches:
                # Get the context before the blank line
                context_start = max(0, match.start() - 50)
                context = text[context_start:match.start()].strip()
                
                fields['text_entries'].append(FormField(
                    field_type='text_entry',
                    label=context,
                    value='',
                    position=position
                ))
            
            position += 1
    
        return fields

    def get_field_values(self, filename: str) -> Dict[str, str]:
        """
        Get a simplified dictionary of field labels and their values.
        
        Args:
            filename: Path to the Word document
            
        Returns:
            Dictionary mapping field labels to their values
        """
        fields = self.extract_fields(filename)
        result = {}
        
        # Process checkboxes
        for checkbox in fields['checkboxes']:
            result[checkbox.label] = checkbox.value
            
        # Process text entries
        for entry in fields['text_entries']:
            result[entry.label] = entry.value
            
        return result

    def get_specific_fields(self, filename: str, field_identifiers: List[str]) -> Dict[str, str]:
        fields = self.extract_fields(filename)
        result = {}
        
        for identifier in field_identifiers:
            # Search through checkboxes
            for checkbox in fields['checkboxes']:
                if identifier.lower() in checkbox.label.lower():
                    result[identifier] = checkbox.value
                    
            # Search through text entries
            for entry in fields['text_entries']:
                if identifier.lower() in entry.label.lower():
                    result[identifier] = entry.value
                    
        return result

def process_legal_form(filename: str):
    print('test')

    parser = LegalFormParser()
    all_fields = parser.extract_fields(filename)

    for checkbox in all_fields['checkboxes']:
        print(f"Checkbox: {checkbox.label} - {checkbox.value}")

    # Get specific fields
    specific_fields = parser.get_specific_fields(filename, [
        "Judicial District",
        "Case Number",
        "Judge",
        "Commissioner",
        "Petitioner",
        "Respondent"
    ])

    # Print all text entry fields
    for entry in specific_fields.items():
        print(entry)
    
    return all_fields, specific_fields