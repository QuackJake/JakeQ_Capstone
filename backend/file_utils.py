import os, json, docx, re, sys
from lxml import etree
from docx import Document
from typing import Union, List, Dict
from pypdf.errors import EmptyFileError
from dataclasses import dataclass
from pypdf import PdfReader, PdfWriter
from datetime import datetime

class docx_reader:
    def __init__(self, target_dir: str, min_underscores=3):
        self.target_dir = target_dir
        if not os.path.exists(target_dir):
            raise ValueError(f"Target directory does not exist: {target_dir}")
        
        self.min_underscores = min_underscores
        
    def exception_wrapper(wrapped_function):
        def _wrapper(*args, **kwargs):
            try:
                result = wrapped_function(*args, **kwargs)
                return result
            except FileNotFoundError as e:
                print(f"File error: {e}")
            except EmptyFileError as e:
                print(f"Empty file error: {e}")
            except Exception as e:
                print(f"An error occurred: {e}")
        return _wrapper
    
    def get_xml(self, file_name, write: bool = False):
        doc = Document(file_name)
        fullText = []
        print(doc)
        
        docx_xml = doc._element.xml
        root = etree.fromstring(docx_xml)
        tree = etree.ElementTree(root)
        
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'v': 'urn:schemas-microsoft-com:vml',
            'w14': 'http://schemas.microsoft.com/office/word/2010/wordml'
        }
        
        # print(tree)
        
        if write:
            file_name_no_ext = os.path.splitext(file_name)[0]
            tree.write((file_name_no_ext + '.xml'), encoding='UTF-8', xml_declaration=True)
        
    
    @exception_wrapper
    def get_text(self, filename, write: bool = False):
        doc = Document(filename)
        fullText = []
        
        print(doc.tables)
    
        def clean_text(text):
            if text is None:
                return ""
            # Replace problematic characters or normalize them
            text = text.replace('\u25ba', '->') # Replace right-pointing triangle
            text = text.replace('\u2022', '*')  # Replace bullet point
            text = text.replace('\u2013', '-')  # Replace en dash
            text = text.replace('\u2014', '--') # Replace em dash
            text = text.replace('\u2018', "'")  # Replace left single quote
            text = text.replace('\u2019', "'")  # Replace right single quote
            text = text.replace('\u201c', '"')  # Replace left double quote
            text = text.replace('\u201d', '"')  # Replace right double quote
            return text.strip()
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        fullText.append(cell.text)

        for para in doc.paragraphs:
            if para.text.strip():
                fullText.append(clean_text(para.text))
        
        docx_xml = doc._element.xml
        root = etree.fromstring(docx_xml)
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'v': 'urn:schemas-microsoft-com:vml',
            'w14': 'http://schemas.microsoft.com/office/word/2010/wordml'
        }
        
        # Get content controls (structured document tags)
        for sdt in root.findall('.//w:sdt', namespaces):
            try:
                # Get the tag name/title if available
                tag = sdt.find('.//w:tag', namespaces)
                tag_val = tag.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') if tag is not None else ''
                
                # Get all text elements within the control
                text_elements = sdt.findall('.//w:t', namespaces)
                control_text = ''.join([clean_text(el.text) for el in text_elements if el.text])
                
                if control_text:
                    if tag_val:
                        fullText.append(f"{clean_text(tag_val)}: {control_text}")
                    else:
                        fullText.append(control_text)
            except Exception as e:
                print(f"Error processing SDT: {str(e)}", file=sys.stderr)
        
        # Get simple fields
        for field in root.findall('.//w:fldSimple', namespaces):
            try:
                # Get the field instruction
                instr = field.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instr', '')
                field_text = field.find('.//w:t', namespaces)
                
                if field_text is not None and field_text.text:
                    if instr:
                        fullText.append(f"{clean_text(instr)}: {clean_text(field_text.text)}")
                    else:
                        fullText.append(clean_text(field_text.text))
            except Exception as e:
                print(f"Error processing simple field: {str(e)}", file=sys.stderr)
        
        # Get complex fields
        for field_char in root.findall('.//w:fldChar', namespaces):
            try:
                fld_char_type = field_char.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType')
                if fld_char_type == 'begin':
                    # Find the corresponding instruction text
                    instr_text = field_char.getparent().getnext()
                    if instr_text is not None:
                        instr = instr_text.find('.//w:instrText', namespaces)
                        if instr is not None and instr.text:
                            fullText.append(clean_text(instr.text))
            except Exception as e:
                print(f"Error processing complex field: {str(e)}", file=sys.stderr)
        
        # Get text boxes (shapes)
        for text_box in root.findall('.//v:textbox', namespaces):
            try:
                text_elements = text_box.findall('.//w:t', namespaces)
                text_box_content = ''.join([clean_text(el.text) for el in text_elements if el.text])
                if text_box_content:
                    fullText.append(text_box_content)
            except Exception as e:
                print(f"Error processing text box: {str(e)}", file=sys.stderr)
        
        # Filter out empty strings and join
        filtered_text = list(filter(None, fullText))

        if write:
            try:
                self.write(filtered_text)
            except Exception as e:
                print(f"Error writing output: {str(e)}", file=sys.stderr)

        try:
            for line in filtered_text:  
                print(line + '\n')
            return '\n'.join(filtered_text)
        except UnicodeEncodeError as e:
            return '\n'.join(filtered_text).encode('utf-8', errors='replace').decode('utf-8')
        
    
    @exception_wrapper
    def write(self, input: list[str], filename: str = None):
        if not filename:
            filename = 'default.txt' # Default file write name
        with open(filename, 'w') as t:
            for line in input:
                t.write(line + '\n')
        msg = filename + ' successfully written'
        print(msg)

    @exception_wrapper
    def get_files(self):
        files = os.listdir(self.target_dir)
        return files
    
    @exception_wrapper
    def format_docx(self):
        # formatted_text, fulltext = self.get_text(self.target_dir + '1353FA_Certificate_of_Service_of_Financial_Declaration.docx')
        # self.write(fulltext)
        # print(fulltext)
        
        dir = self.target_dir + '1353FA_Certificate_of_Service_of_Financial_Declaration.docx'
        
        parseda, parseds, c = self.process_legal_form(dir)
        print(parseda)        

    @exception_wrapper
    def extract_metadata(self, filename: str):
        doc = docx.Document(filename)
        core_properties = doc.core_properties
        metadata = {}
        for prop in dir(core_properties):
            if prop.startswith('__'):
                continue
            value = getattr(core_properties, prop)
            if callable(value):
                continue
            if prop == 'created' or prop == 'modified' or prop == 'last_printed':
                if value:
                    value = value.strftime('%Y-%m-%d %H:%M:%S')
                else:
                    value = None
            metadata[prop] = value
        try:
            custom_properties = core_properties.custom_properties
            if custom_properties:
                metadata['custom_properties'] = {}
                for prop in custom_properties:
                    metadata['custom_properties'][prop.name] = prop.value
        except AttributeError:
            pass

        for x in metadata:
            print (x)
            for y in metadata[x]:
                print (y, ':', metadata[x][y])

        
        return metadata

    def xml_to_docx(dumps_dir, output_docx_path, filename='1353FA_Certificate_of_Service_of_Financial_Declaration.xml'):
        # Construct the full XML file path
        xml_file_path = os.path.join(dumps_dir, filename)
        
        # Check if the file exists
        if not os.path.exists(xml_file_path):
            print(f"File not found: {xml_file_path}")
            return
        
        # Parse the XML file
        tree = etree.parse(xml_file_path)
        root = tree.getroot()
        
        # Define the namespace mapping for the 'w' prefix
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        # Find the <w:body> element using the namespace dictionary
        body = root.find('w:body', namespaces=ns)
        if body is None:
            print("No <w:body> element found in the XML.")
            return
        
        # Create a new DOCX Document
        doc = Document()
        
        # Iterate over each <w:p> element (paragraph) inside the <w:body>
        for p in body.findall('w:p', namespaces=ns):
            # DOCX XML often splits text across multiple <w:t> tags inside a paragraph.
            # We collect all <w:t> elements to reconstruct the paragraph text.
            texts = [node.text for node in p.findall('.//w:t', namespaces=ns) if node.text]
            full_text = ''.join(texts)
            doc.add_paragraph(full_text)
        
        # Save the DOCX file to the specified path
        doc.save(output_docx_path)
        print(f"DOCX file successfully saved to {output_docx_path}")

    
    def replace_fields_with_placeholders(self, doc_path, output_path):
        doc = Document(doc_path)

        # Define patterns for detection
        patterns = {
            r'_____': '$Line',  # Replace exactly five underscores
            r'\[  ?\]': '[$C]',  # Checkboxes [ ]
        }

        # Process paragraphs in the main body
        self.process_paragraphs(doc.paragraphs, patterns)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.process_paragraphs(cell.paragraphs, patterns)

        # Process headers & footers
        for section in doc.sections:
            self.process_paragraphs(section.header.paragraphs, patterns)
            self.process_paragraphs(section.footer.paragraphs, patterns)

        # Ensure the output directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        doc.save(output_path)
        print(f"Processed document saved as: {output_path}")

    def process_paragraphs(self, paragraphs, patterns):
        """Apply regex replacements to a list of paragraphs."""
        for para in paragraphs:
            for pattern, placeholder in patterns.items():
                para.text = re.sub(pattern, placeholder, para.text)
    

# @dataclass
# class FormField:
#     field_type: str  # 'checkbox' or 'text_entry'
#     label: str       # The text associated with the field
#     value: str       # The current value or empty string
#     position: int    # Position in document for reference

# class LegalFormParser:
#     def __init__(self):
#         self.checkbox_pattern = r'\[\s*[Xx]?\s*\]'  # Matches [  ] or [X]
#         self.blank_line_pattern = r'_{3,}'          # Matches 3 or more underscores
        
#     def extract_fields(self, filename: str) -> Dict[str, List[FormField]]:
#         """
#         Extract form fields from a legal document.
        
#         Args:
#             filename: Path to the Word document
            
#         Returns:
#             Dictionary containing lists of checkboxes and text entry fields
#         """
#         doc = Document(filename)
#         fields = {
#             'checkboxes': [],
#             'text_entries': []
#         }
        
#         position = 0
        
#         for paragraph in doc.paragraphs:
#             text = paragraph.text.strip()
#             if not text:
#                 continue
                
#             # Find checkbox fields
#             checkbox_matches = list(re.finditer(self.checkbox_pattern, text))
#             for match in checkbox_matches:
#                 # Get the label by looking at text after the checkbox
#                 start_pos = match.end()
#                 # Find the next checkbox or end of line
#                 next_checkbox = re.search(self.checkbox_pattern, text[start_pos:])
#                 end_pos = start_pos + next_checkbox.start() if next_checkbox else len(text)
#                 label = text[start_pos:end_pos].strip()
                
#                 fields['checkboxes'].append(FormField(
#                     field_type='checkbox',
#                     label=label,
#                     value='checked' if 'X' in match.group() else 'unchecked',
#                     position=position
#                 ))
            
#             # Find text entry fields
#             blank_matches = list(re.finditer(self.blank_line_pattern, text))
#             for match in blank_matches:
#                 # Get the context before the blank line
#                 context_start = max(0, match.start() - 50)
#                 context = text[context_start:match.start()].strip()
                
#                 fields['text_entries'].append(FormField(
#                     field_type='text_entry',
#                     label=context,
#                     value='',
#                     position=position
#                 ))
            
#             position += 1
    
#         return fields

#     def get_field_values(self, filename: str) -> Dict[str, str]:
#         """
#         Get a simplified dictionary of field labels and their values.
        
#         Args:
#             filename: Path to the Word document
            
#         Returns:
#             Dictionary mapping field labels to their values
#         """
#         fields = self.extract_fields(filename)
#         result = {}
        
#         # Process checkboxes
#         for checkbox in fields['checkboxes']:
#             result[checkbox.label] = checkbox.value
            
#         # Process text entries
#         for entry in fields['text_entries']:
#             result[entry.label] = entry.value
            
#         return result

#     def get_specific_fields(self, filename: str, field_identifiers: List[str]) -> Dict[str, str]:
#         fields = self.extract_fields(filename)
#         result = {}
        
#         for identifier in field_identifiers:
#             # Search through checkboxes
#             for checkbox in fields['checkboxes']:
#                 if identifier.lower() in checkbox.label.lower():
#                     result[identifier] = checkbox.value
                    
#             # Search through text entries
#             for entry in fields['text_entries']:
#                 if identifier.lower() in entry.label.lower():
#                     result[identifier] = entry.value
                    
#         return result

# def process_legal_form(filename: str):
    # print('test')

    # parser = LegalFormParser()
    # all_fields = parser.extract_fields(filename)

    # for checkbox in all_fields['checkboxes']:
    #     print(f"Checkbox: {checkbox.label} - {checkbox.value}")

    # # Get specific fields
    # specific_fields = parser.get_specific_fields(filename, [
    #     "Judicial District",
    #     "Case Number",
    #     "Judge",
    #     "Commissioner",
    #     "Petitioner",
    #     "Respondent"
    # ])

    # # Print all text entry fields
    # for entry in specific_fields.items():
    #     print(entry)
    
    # return all_fields, specific_fields